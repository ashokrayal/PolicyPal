"""
Script to create test files for DOCX and CSV formats.
"""
from pathlib import Path
import pandas as pd
from docx import Document

def create_test_docx():
    """Create a simple test DOCX file."""
    test_dir = Path("tests/test_data")
    test_dir.mkdir(parents=True, exist_ok=True)
    
    docx_path = test_dir / "test_document.docx"
    
    # Create DOCX document
    doc = Document()
    
    # Add title
    doc.add_heading('Test Policy Document', 0)
    
    # Add content
    doc.add_paragraph('This is a test policy document for the PolicyPal project.')
    doc.add_paragraph('')
    
    doc.add_heading('Section 1: Introduction', level=1)
    doc.add_paragraph('This document contains sample policy information that can be used to test the DOCX parsing functionality of our RAG chatbot.')
    doc.add_paragraph('')
    
    doc.add_heading('Section 2: Vacation Policy', level=1)
    doc.add_paragraph('Employees are entitled to 20 vacation days per year.')
    doc.add_paragraph('Vacation requests must be submitted at least 2 weeks in advance.')
    doc.add_paragraph('')
    
    doc.add_heading('Section 3: Remote Work Policy', level=1)
    doc.add_paragraph('Employees may work remotely up to 3 days per week.')
    doc.add_paragraph('Remote work requires manager approval and stable internet connection.')
    
    # Save document
    doc.save(str(docx_path))
    print(f"Test DOCX created at: {docx_path}")
    return docx_path

def create_test_csv():
    """Create a simple test CSV file."""
    test_dir = Path("tests/test_data")
    test_dir.mkdir(parents=True, exist_ok=True)
    
    csv_path = test_dir / "test_data.csv"
    
    # Create sample data
    data = {
        'Policy_Name': ['Vacation Policy', 'Remote Work Policy', 'Dress Code Policy', 'Expense Policy'],
        'Category': ['Time Off', 'Work Arrangement', 'Appearance', 'Finance'],
        'Description': [
            'Employees get 20 vacation days per year',
            'Up to 3 remote work days per week allowed',
            'Business casual dress code required',
            'Expense reports due within 30 days'
        ],
        'Effective_Date': ['2024-01-01', '2024-01-01', '2024-01-01', '2024-01-01'],
        'Status': ['Active', 'Active', 'Active', 'Active']
    }
    
    # Create DataFrame and save to CSV
    df = pd.DataFrame(data)
    df.to_csv(csv_path, index=False)
    print(f"Test CSV created at: {csv_path}")
    return csv_path

def create_all_test_files():
    """Create all test files."""
    print("Creating test files...")
    create_test_docx()
    create_test_csv()
    print("All test files created successfully!")

if __name__ == "__main__":
    create_all_test_files() 